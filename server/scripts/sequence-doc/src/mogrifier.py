'''
This takes the source document and applies selections. Using logic baked into
the source document it gets modified based on the selections and the final document
gets returned, with all template language removed
'''
import re
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
import logging
import utils
from typing import Dict, List

logging.getLogger().setLevel(logging.DEBUG)

P_TAG = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
BOOKMARK_TAGS = ['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkEnd', "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart"]
SECTION_TAG = ["{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"]
TABLE_TAG = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl'
ANNOTATION_STYLE = 'Toggle'

OP_LIST = ['AND', 'OR', 'YES', 'NO', 'EQUALS', 'NOT_EQUALS', 'ANY', 'DELETE']
TABLE_OP_LIST = ['TABLE', 'ROW', 'COLUMN']
# Type hints
Selections = Dict[str, List]

elements_to_delete = []

def initialize_remove_list():
    ''' Resets global remove list to an empty array
    '''
    global elements_to_delete
    elements_to_delete = []

def remove_node(node):
    ''' Controls access to a global list for nodes to be deleted
    '''
    try:
       remove_element(node._element)
    except AttributeError:
        # node is actually an element, not a Paragraph
        remove_element(node)

def remove_element(element):
    ''' Adds a lxml element to global list for deletion
    '''
    global elements_to_delete
    elements_to_delete.append(element)

def get_heading_level(paragraph):
    ''' Gets heading level of paragraph
    '''
    style = paragraph.style.name
    match = re.match(r'Heading (\d+)', style)

    if not match:
        logging.info('Saw unrecognized style: {}'.format(style))
        return 100
    return int(match.group(1))

def remove_info_box(paragraph, run_op_lookup: Dict):
    '''
        Removes info box text within a section (paragraph is a detected info box in a section)
        Also removes any sibling elements within the paragraph
    '''
    remove_node(paragraph)

    for sib_el in paragraph._element.itersiblings(P_TAG):
        sib_p = Paragraph(sib_el, paragraph._parent)
        sib_level = get_heading_level(sib_p)
        
        if (sib_level == 0):
            for run in sib_p.runs:
                if run.element in run_op_lookup and run_op_lookup[run.element]['op'] in ('VENT', 'ENERGY'):
                    # We've reached another VENT/ENERGY toggle, stop
                    return
        else:
            # reached end of info box, stop
            return
                
        remove_node(sib_p)

def remove_section(paragraph: Paragraph, run_op_lookup: Dict):
    ''' Removes a section of text and its siblings
    '''
    style = paragraph.style.name
    level = get_heading_level(paragraph)
    
    if style in ('Info. box', 'InfoboxList'):
        return remove_info_box(paragraph, run_op_lookup)

    for sib_el in paragraph._element.itersiblings():
        if sib_el.tag == P_TAG:
            sib_p = Paragraph(sib_el, paragraph._parent)
            sib_level = get_heading_level(sib_p)

            if (sib_level <= level) and (sib_level != 9999):
                break

            remove_node(sib_p)
        elif sib_el.tag == TABLE_TAG:
            # Assume table matches indentation of parent paragraph
            # so just go ahead and remove
            sib_t = Table(sib_el, paragraph._parent)
            remove_node(sib_t)
        elif sib_el.tag in BOOKMARK_TAGS:
            remove_node(sib_el)
        elif sib_el.tag in SECTION_TAG:
            break
        else:
            logging.error('Saw unrecognized tag "%s"', sib_el.tag)

    remove_node(paragraph)

def edit_table(table_item):
    ''' Removes items from table or whole table depending on the operation
    '''
    if table_item['op'] == 'COLUMN':
        for cell in table_item['column'].cells:
            remove_node(cell)
    else:
        remove_node(table_item)

def get_column(table: Table, table_cell):
    ''' Determins the column of a cell within a table
    '''
    for column in table.columns:
        for cell in column.cells:
            if cell.text == table_cell.text:
                return column

def create_control_structures(doc):
    ''' Sets up the control structure for Sections and Tables
    '''
    control_structures = []
    run_op_lookup = {}

    current = None

    ## Sections
    for para in doc.paragraphs:
        for run in para.runs:
            if (
                (current and run.style.name != ANNOTATION_STYLE)
                or
                (current and para is not current['paragraph'])
            ):
                current['text'] = current['text'].strip()
                if current['text']:
                    control_structures.append(current)
                current = None
                
            if not current and run.style.name == ANNOTATION_STYLE:
                current = {
                    'paragraph': para,
                    'text': '',
                    'runs': [],
                }
                run_op_lookup[run.element] = current

            if current and run.style.name == ANNOTATION_STYLE:
                run_op_lookup[run.element] = current
                current['runs'].append(run)
                current['text'] += run.text

    ## Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if (
                            (current and run.style.name != ANNOTATION_STYLE)
                            or
                            (current and para is not current['paragraph'])
                        ):
                            current['text'] = current['text'].strip()
                            if current['text']:
                                control_structures.append(current)
                            current = None

                        if not current and run.style.name == ANNOTATION_STYLE:
                            current = {
                                'paragraph': para,
                                'text': '',
                                'runs': [],
                                'table': table,
                                'row': row,
                                'column': get_column(table, cell),
                                'cell': cell,
                            }
                            run.op = current

                        if current and run.style.name == ANNOTATION_STYLE:
                            run.op = current
                            current['runs'].append(run)
                            current['text'] += run.text

    for op in control_structures:
        tokens = utils.remove_empty_strings(re.split(r'\W+', op['text']))
        if tokens:
            op['op'] = tokens[0]
        else:
            op['op'] = ''

    # Debugging
    # for op in control_structures:
    #     logging.info(op)

    return control_structures, run_op_lookup

def remove_info_and_instr_boxes(doc, selections: Selections):
    ''' Removes info and instruction boxes
    '''
    info_box_style = 'Info. box'
    instr_box_style = 'Instr. box'

    for para in doc.paragraphs:
        if utils.reduce_to_boolean(selections['DEL_INFO_BOX']) and para.style.name == info_box_style:
            remove_node(para)
        if para.style.name == instr_box_style:
            remove_node(para)

def evaluate_annotation(op, name_map, selections: Selections):
    '''
        Determines the operation and decides if the operation requires us to delete (return True).
        This is a recursive funtion if the operation is an AND or OR.
        Tables use this function as well after their table operation logic is determined.

        Example toggles:
        + `[YES have_CO2Sen]` – Keep this section if a CO2 sensor is present. Otherwise remove.
        + `[NO have_CO2Sen]` – Keep this section if a CO2 sensor is not present. Otherwise remove.
        + `[EQUALS buiPreCon ReliefFan]` – Keep this section if the type of building pressure control system is an actuated relief damper, with relief fan(s). Otherwise remove.
        + `[NOT_EQUALS buiPreCon RelifFan]`  Keep this section if the type of building pressure control system is *not* an actuated relief damper, with relief fans(s). Otherwise remove.
        + `[ANY buiPreCon ReturnFanMeasuredAir ReturnFanCalculatedAir]` – Keep this section if the type of building pressure control system is a return fan, tracking measured supply and return airflow or a return fan, tracking calculated supply and return airflow. Otherwise remove.
        + `[AND [any toggle] [any toggle]]` - Keep this section if both nested toggles would keep the section. The nested toggles can be any of these toggles including AND or OR. Otherwise remove.
        + `[OR [any toggle] [any toggle]]` - Keep this section if one of the nested toggles would keep the section. The nested toggles can be any of these toggles including AND or OR. Otherwise remove.
        + `[DELETE]` - Remove this section.
    '''
    if op['op'] == 'YES':
        # 1: map short name to long name
        tokens = utils.remove_empty_strings(re.split(r'\W+', op['text']))
        if len(tokens) == 2:
            short_name = tokens[1]
        else:
            logging.error('Invalid operation: %s', op['text'])
            return False

        if short_name not in name_map:
            logging.error('%s not found', short_name)
            return False
            
        long_name = name_map[short_name]

        # 2: check if relevant selection is available
        if long_name not in selections:
            logging.error('Path "%s" not found in store, deleting', long_name)
            return True
        # 3: apply the operation type
        elif not utils.reduce_to_boolean(selections[long_name]):
            return True
            
    if op['op'] == 'NO':
        tokens = utils.remove_empty_strings(re.split(r'\W+', op['text']))
        if len(tokens) == 2:
            short_name = tokens[1]
        else:
            logging.error('Invalid operation: %s', op['text'])
            return False

        if short_name not in name_map:
            logging.error('%s not found', short_name)
            return False
            
        long_name = name_map[short_name]

        if long_name not in selections:
            logging.error('Path "%s" not found in store, deleting', long_name)
            return True
        elif utils.reduce_to_boolean(selections[long_name]):
            return True

    if op['op'] == 'EQUALS':
        tokens = utils.remove_empty_strings(re.split(r'\W+', op['text']))
        if len(tokens) == 3:
            short_name = tokens[1]
            short_compare = tokens[2]
        else:
            logging.error('Invalid operation: %s', op['text'])
            return False
        
        if short_name not in name_map:
            logging.error('%s not found', short_name)
            return False

        if short_compare not in name_map:
            logging.error('%s not found', short_compare)
            return False

        long_name = name_map[short_name]
        long_compare = name_map[short_compare]
                
        if long_name not in selections:
            logging.error('Path "%s" not found in store, deleting', long_name)
            return True
        elif long_compare not in selections[long_name]:
            return True

    if op['op'] == 'NOT_EQUALS':
        tokens = utils.remove_empty_strings(re.split(r'\W+', op['text']))
        if len(tokens) == 3:
            short_name = tokens[1]
            short_compare = tokens[2]
        else:
            logging.error('Invalid operation: %s', op['text'])
            return False
        
        if short_name not in name_map:
            logging.error('%s not found', short_name)
            return False

        if short_compare not in name_map:
            logging.error('%s not found', short_compare)
            return False

        long_name = name_map[short_name]
        long_compare = name_map[short_compare]
                
        if long_name not in selections:
            logging.error('Path "%s" not found in store, deleting', long_name)
            return True
        elif long_compare in selections[long_name]:
            return True
            
    if op['op'] == 'ANY':
        tokens = [token for token in re.split(r'\W+', op['text']) if token]
        if len(tokens) >= 3:
            short_name = tokens[1]
            short_compare = tokens[2:]
        else:
            logging.error('Invalid operation: %s', op['text'])
            return False

        if short_name not in name_map:
            logging.error('%s not found', short_name)
            return False
            
        long_name = name_map[short_name]
        long_compare = map(lambda name: name_map[name], short_compare)
                
        if long_name not in selections:
            logging.error('Path "%s" not found in store, deleting', long_name)
            return True
        elif not utils.common_member(selections[long_name], long_compare):
            return True
            
    if op['op'] == 'DELETE':
        return True

    if op['op'] == 'AND':
        match = re.match(r'\[AND \[(.+)] \[(.+)]]', op['text'])
        if not match:
            logging.error('Invalid format for tag:  %s', op['text'])
            return False

        condition_one = condition_two = False
        group_one_tokens = re.split(r'\W+', match.group(1))
        group_one_op = group_one_tokens[0]
        group_two_tokens = re.split(r'\W+', match.group(2))
        group_two_op = group_two_tokens[0]

        if group_one_op in OP_LIST:
            nested_op = {**op, 'text': f'[{match.group(1)}]', 'op': group_one_op}
            condition_one = evaluate_annotation(nested_op, name_map, selections)
        else:
            logging.error('Unknown operation: %s', group_one_op)
            return False

        if group_two_op in OP_LIST:
            nested_op = {**op, 'text': f'[{match.group(2)}]', 'op': group_two_op}
            condition_two = evaluate_annotation(nested_op, name_map, selections)
        else:
            logging.error('Unknown operation: %s', group_two_op)
            return False

        return condition_one and condition_two

    if op['op'] == 'OR':
        match = re.match(r'\[OR \[(.+)] \[(.+)]]', op['text'])
        if not match:
            logging.error('Invalid format for tag:  %s', op['text'])
            return False

        condition_one = condition_two = False
        group_one_tokens = re.split(r'\W+', match.group(1))
        group_one_op = group_one_tokens[0]
        group_two_tokens = re.split(r'\W+', match.group(2))
        group_two_op = group_two_tokens[0]

        if group_one_op in OP_LIST:
            nested_op = {**op, 'text': f'[{match.group(1)}]', 'op': group_one_op}
            condition_one = evaluate_annotation(nested_op, name_map, selections)
        else:
            logging.error('Unknown operation: %s', group_one_op)
            return False

        if group_two_op in OP_LIST:
            nested_op = {**op, 'text': f'[{match.group(2)}]', 'op': group_two_op}
            condition_two = evaluate_annotation(nested_op, name_map, selections)
        else:
            logging.error('Unknown operation: %s', group_two_op)
            return False

        return condition_one or condition_two

    return False

def apply_selections(control_structure, name_map, run_op_lookup, selections: Selections):
    ''' 
        Determines how to handle toggles, section vs table

        Example section toggles:
        + `[YES have_CO2Sen]` – Keep this section if a CO2 sensor is present. Otherwise remove.
        + `[NO have_CO2Sen]` – Keep this section if a CO2 sensor is not present. Otherwise remove.
        + `[EQUALS buiPreCon ReliefFan]` – Keep this section if the type of building pressure control system is an actuated relief damper, with relief fan(s). Otherwise remove.
        + `[NOT_EQUALS buiPreCon RelifFan]` - Keep this section if the type of building pressure control system is *not* an actuated relief damper, with relief fans(s). Otherwise remove.
        + `[ANY buiPreCon ReturnFanMeasuredAir ReturnFanCalculatedAir]` – Keep this section if the type of building pressure control system is a return fan, tracking measured supply and return airflow or a return fan, tracking calculated supply and return airflow. Otherwise remove.
        + `[AND [any toggle] [any toggle]]` - Keep this section if both nested toggles would keep the section. The nested toggles can be any of these toggles including AND or OR.  Otherwise remove.
        + `[OR [any toggle] [any toggle]]` - Keep this section if one of the nested toggles would keep the section. The nested toggles can be any of these toggles including AND or OR.  Otherwise remove.
        + `[DELETE]` - Remove this section.

        Example table toggles:
        + `[TABLE YES have_CO2Sen]` - Keep the whole table if a CO2 sensor is present. Otherwise remove.
        + `[ROW YES have_CO2Sen]` - Keep this row if a CO2 sensor is present. Otherwise remove.
        + `[COLUMN YES have_CO2Sen]` - Keep this column if a CO2 sensor is present. Otherwise remove.
        + `[COLUMN AND [any toggle] [any toggle]]` - Keep this column if both nested toggles would keep the column. Otherwise remove.
    '''
    for op in control_structure:
        if op['op'] in OP_LIST:
            if evaluate_annotation(op, name_map, selections):
                remove_section(op['paragraph'], run_op_lookup)

        if op['op'] in TABLE_OP_LIST:
            text_op = op['text'].replace(f"{op['op']} ", "")
            tokens = utils.remove_empty_strings(re.split(r'\W+', text_op))
            next_op = tokens[0]

            if next_op in OP_LIST:
                table_op = {**op, 'text': text_op, 'op': next_op}
                if evaluate_annotation(table_op, name_map, selections):
                    edit_table(op)
            else:
                logging.error('Invlaid format for tag: %s', op['text'])


    # return not necessary - just reinforcing that control_structure is what is modified
    return control_structure

def convert_units(control_structure, name_map, selections: Selections):
    ''' Based on unit selection, goes through doc to update
    '''
    short_name = "UNITS"
    si_short_name = "SI"
    ip_short_name = "IP"

    if short_name not in name_map:
        logging.error('%s not found', short_name)
        return
    if si_short_name not in name_map:
        logging.error('%s not found', si_short_name)
        return
    if ip_short_name not in name_map:
        logging.error('%s not found', ip_short_name)
        return

    for op in control_structure:
        if op['op'] == short_name:
            match = re.match(r'\[UNITS \[(.+)] \[(.+)]]', op['text'])
            if not match:
                logging.error('Invalid format for tag:  %s', op['text'])
                continue
            long_name = name_map[short_name]
            unit_selection = selections[long_name][0]
            if unit_selection == name_map[si_short_name]:
                # TODO: use this as an approach for 'writes' to the docx 
                op['runs'][0].text = match.group(1)
                op['runs'][0].style = None
            elif unit_selection == name_map[ip_short_name]:
                op['runs'][0].text = match.group(2)
                op['runs'][0].style = None
            else:
                logging.error('"%s" is not a valid unit system', unit_selection)

def remove_toggles(doc):
    ''' Step through and remove 'toggle' text
    '''
    for para in doc.paragraphs:
        for run in para.runs:
            if run.style.name == ANNOTATION_STYLE:
                remove_node(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.style.name == ANNOTATION_STYLE:
                            remove_node(run)

def mogrify_doc(doc: Document, name_map: Dict, selections: Selections) -> Document:
    ''' Applies selections to the provided document. This mutates the provided
        document
    '''
    initialize_remove_list()
    # walk through source_doc to find each conditional point in the doc
    control_structure, run_op_lookup = create_control_structures(doc)

    # Remove info and Instruction Boxes - updates remove_list internally
    remove_info_and_instr_boxes(doc, selections)

    # apply all paragraph and table selections
    apply_selections(control_structure, name_map, run_op_lookup, selections)

    # convert units
    convert_units(control_structure, name_map, selections)

    # remove toggle text
    remove_toggles(doc)

    # finally remove all nodes flagged for deletion
    for el in elements_to_delete:
        try:
            el.getparent().remove(el)
        except AttributeError:
            logging.info('Element "%s" was already gone', el)
            pass

    return doc
